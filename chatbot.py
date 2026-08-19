from __future__ import annotations

import os
import re
import gc
import ssl
import glob
import pickle
import logging
import unicodedata
import subprocess
from datetime import datetime
from typing import Optional

import streamlit as st
from docx import Document
import pdfplumber

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import Docx2txtLoader, PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain_community.retrievers import BM25Retriever
from langchain_core.documents import Document as LCDocument
from langchain_huggingface import HuggingFaceEmbeddings

ssl._create_default_https_context = ssl._create_unverified_context

# --------------------------------------------------------------------------- #
# Configuración global
# --------------------------------------------------------------------------- #

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger("bot_resoluciones")

CHUNK_SIZE = 800
CHUNK_OVERLAP = 200
TOP_K_INITIAL = 30
TOP_K_FINAL = 12
MAX_HISTORY_TURNS = 3
RECENCY_BOOST = 0.6
CURRENT_YEAR = datetime.now().year
INDEX_VERSION = "v2-tablas"
BATCH_SIZE_DOCS = 50
GC_EVERY_N_DOCS = 25

N_RESOLUCIONES_EXPANDIR = 3
MAX_CHUNKS_POR_RES_EXPANDIDA = 10
MAX_CONTEXT_CHARS = 14000
MAX_LEN_PREGUNTA_REESCRITA = 300

N_VARIANTES_QUERY = 3
RECENCY_TIEBREAK_NEUTRO = 0.15
MAX_LEN_VARIANTE = 200


def _detect_device() -> str:
    try:
        import torch
        if torch.backends.mps.is_available():
            return "mps"
        if torch.cuda.is_available():
            return "cuda"
    except Exception:  # noqa: BLE001
        pass
    return "cpu"

DEVICE = _detect_device()

# --------------------------------------------------------------------------- #
# Detección de entorno (Ollama, RAM)
# --------------------------------------------------------------------------- #

def _ollama_running() -> bool:
    try:
        out = subprocess.run(
            ["pgrep", "-x", "ollama"], capture_output=True, timeout=2
        )
        return out.returncode == 0
    except Exception:  # noqa: BLE001
        return False

def _ram_disponible_gb() -> Optional[float]:
    try:
        import psutil
        return psutil.virtual_memory().available / (1024 ** 3)
    except Exception:  # noqa: BLE001
        return None

# --------------------------------------------------------------------------- #
# Streamlit
# --------------------------------------------------------------------------- #

st.set_page_config(page_title="Bot de Resoluciones", layout="wide")
st.title("🔍 Bot de Resoluciones")
st.write(
    "Consultá resoluciones y obtené respuestas detalladas con cita de archivos. "
    "El sistema prioriza la información más reciente y señala contradicciones "
    "entre resoluciones cuando las detecta."
)

# --------------------------------------------------------------------------- #
# Patrones de identificación
# --------------------------------------------------------------------------- #

PATRONES = [
    re.compile(r'(?P<letras>[A-Z]{1,5})-(?P<num1>\d{3})[-\.](?P<num2>\d{2,4})', re.IGNORECASE),
    re.compile(r'(?P<letras>[A-Z]{3,5})[-\.](?P<letras2>[A-Z]{2,4})[-\.](?P<num>\d{1,4})', re.IGNORECASE),
    re.compile(r'(?P<letras1>[A-Z]{1})-(?P<letras2>[A-Z]{3})[-\.](?P<num>\d{2,4})', re.IGNORECASE),
    re.compile(r'(?P<letras>[A-Z]{3,5})[-\.](?P<num1>\d{1,4})[-\.](?P<num2>\d{1,4})', re.IGNORECASE),
    re.compile(r'(?P<letras>[A-Z]{3,5})[-\.](?P<num1>\d{1,4})', re.IGNORECASE),
    re.compile(r'(?P<letras>[A-Z]{3,5})\s+(?P<num1>\d{3})[-\.](?P<num2>\d{2,4})', re.IGNORECASE),
    re.compile(r'(?P<letras>[A-Z]{3,5})-?\s*(?P<num1>\d{3})[-\.]{1,2}(?P<num2>\d{2,4})', re.IGNORECASE),
    re.compile(r'(?P<letras>[A-Z]{1,2})-(?P<num1>\d{3})[-\.](?P<num2>\d{2})', re.IGNORECASE),
    re.compile(r'(?P<letras>[A-Z]{3,5})-\.(?P<num1>\d{2,4})', re.IGNORECASE),
]

PATRON_ANIO_4 = re.compile(r'(19\d{2}|20\d{2})')
PATRON_ANIO_2 = re.compile(r'[-\.\s/](\d{2})(?!\d)')

def extraer_identificador(nombre_archivo: str) -> str:
    for p in PATRONES:
        m = p.search(nombre_archivo)
        if m:
            return m.group(0)
    return nombre_archivo

def extraer_anio(nombre: str, contenido: str = "") -> int:
    """Devuelve año entero o 0 si no se encuentra."""
    m4 = PATRON_ANIO_4.search(nombre)
    if m4:
        return int(m4.group(1))
    m2 = PATRON_ANIO_2.search(nombre)
    if m2:
        n = int(m2.group(1))
        return 1900 + n if n > 50 else 2000 + n
    if contenido:
        m4c = PATRON_ANIO_4.search(contenido[:500])
        if m4c:
            return int(m4c.group(1))
    return 0

# --------------------------------------------------------------------------- #
# Carga de un solo documento
# --------------------------------------------------------------------------- #

HEADER_HINTS = {
    "asignatura", "materia", "cátedra", "catedra", "comisión", "comision",
    "profesor", "profesora", "profesor/a", "titular", "adjunto", "adjunta",
    "asistente", "auxiliar", "ayudante", "jtp", "jefe", "encargado",
    "cargo", "rol", "función", "funcion", "dedicación", "dedicacion",
    "apellido", "nombre", "docente", "responsable", "carrera", "año",
    "cuatrimestre", "código", "codigo", "área", "area", "departamento",
}

def _fila_es_basura(celdas: list[str]) -> bool:
    no_vacias = [c for c in celdas if c]
    if not no_vacias:
        return True
    if all(len(set(c.replace(" ", ""))) <= 1 for c in no_vacias):
        return True
    return False

def _es_fila_header(celdas: list[str]) -> bool:
    no_vacias = [c.strip().lower() for c in celdas if c and c.strip()]
    if not no_vacias:
        return False
    aciertos = sum(
        1 for celda in no_vacias
        if any(hint in celda for hint in HEADER_HINTS)
    )
    return aciertos >= 2 or (len(no_vacias) >= 2 and aciertos / len(no_vacias) >= 0.4)

def _normalizar_header(celda: str) -> str:
    c = celda.strip()
    c = re.sub(r"\s+", " ", c)
    c = re.sub(r"\s*\(.*?\)\s*", "", c)
    return c or "Columna"

def _emitir_fila_keyvalue(headers: list[str], celdas: list[str]) -> str:
    pares = []
    for i, val in enumerate(celdas):
        if i >= len(headers):
            break
        val = val.strip()
        if not val or val.lower() in {"-", "—", "s/d", "n/a"}:
            continue
        clave = headers[i] if i < len(headers) else f"Col{i+1}"
        pares.append(f"{clave}: {val}")
    return " | ".join(pares) if pares else ""

def _procesar_tabla_generica(filas: list[list[str]], tabla_idx: int) -> list[str]:
    filas_limpias: list[list[str]] = []
    for fila in filas:
        if not fila:
            continue
        celdas = [str(c).strip() if c else "" for c in fila]
        if _fila_es_basura(celdas):
            continue
        filas_limpias.append(celdas)

    if not filas_limpias:
        return []

    header_idx = None
    for i, fila in enumerate(filas_limpias[:3]):
        if _es_fila_header(fila):
            header_idx = i
            break

    salida: list[str] = []
    if header_idx is not None:
        headers = [_normalizar_header(c) for c in filas_limpias[header_idx]]
        salida.append(f"--- Tabla {tabla_idx + 1} (columnas: {', '.join(headers)}) ---")
        for fila in filas_limpias[header_idx + 1:]:
            linea = _emitir_fila_keyvalue(headers, fila)
            if linea:
                salida.append(linea)
    else:
        salida.append(f"--- Tabla {tabla_idx + 1} (sin encabezado detectado) ---")
        for fila in filas_limpias:
            no_vacias = [c for c in fila if c]
            if no_vacias:
                salida.append(" | ".join(no_vacias))

    return salida

def _procesar_tablas_docx(path: str) -> list[str]:
    frases: list[str] = []
    try:
        doc = Document(path)
        for idx, table in enumerate(doc.tables):
            filas = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
            ]
            frases.extend(_procesar_tabla_generica(filas, idx))
        del doc
    except Exception as e:  # noqa: BLE001
        logger.warning("Error tablas docx %s: %s", path, e)
    return frases

def _procesar_tablas_pdf(path: str) -> list[str]:
    frases: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            tabla_global_idx = 0
            for page in pdf.pages:
                try:
                    tablas = page.extract_tables() or []
                    for tabla in tablas:
                        frases.extend(
                            _procesar_tabla_generica(tabla, tabla_global_idx)
                        )
                        tabla_global_idx += 1
                finally:
                    page.flush_cache()
    except Exception as e:  # noqa: BLE001
        logger.warning("Error tablas pdf %s: %s", path, e)
    return frases

def load_document(file_path: str) -> list[LCDocument]:
    try:
        if file_path.endswith(".docx"):
            base = Docx2txtLoader(file_path).load()
            tablas = _procesar_tablas_docx(file_path)
        elif file_path.endswith(".pdf"):
            base = PyPDFLoader(file_path).load()
            tablas = _procesar_tablas_pdf(file_path)
        else:
            return []

        nombre = os.path.basename(file_path)
        identificador = extraer_identificador(nombre)
        primer_contenido = base[0].page_content if base else ""
        anio = extraer_anio(nombre, primer_contenido)

        meta = {
            "resolucion": identificador,
            "source": nombre,
            "anio": anio,
            "path": file_path,
        }
        for d in base:
            d.metadata.update(meta)

        tabla_docs = [
            LCDocument(
                page_content=(
                    f"[Tabla | Resolución: {identificador} | "
                    f"Año: {anio or 's/d'} | Archivo: {nombre}] {f}"
                ),
                metadata=meta,
            )
            for f in tablas
        ]
        return base + tabla_docs
    except Exception as e:  # noqa: BLE001
        logger.error("Error cargando %s: %s", file_path, e)
        return []

# --------------------------------------------------------------------------- #
# Persistencia de estado
# --------------------------------------------------------------------------- #

def _file_signature(path: str) -> tuple[float, int]:
    s = os.stat(path)
    return (s.st_mtime, s.st_size)

def guardar_estado(estado_path: str, estado: dict) -> None:
    with open(estado_path, "wb") as f:
        pickle.dump({"version": INDEX_VERSION, "files": estado}, f)

def cargar_estado(estado_path: str) -> dict:
    if not os.path.exists(estado_path):
        return {}
    try:
        with open(estado_path, "rb") as f:
            data = pickle.load(f)
        if data.get("version") != INDEX_VERSION:
            logger.info("Versión de índice cambió → reindexación completa.")
            return {}
        return data.get("files", {})
    except Exception as e:  # noqa: BLE001
        logger.warning("Estado corrupto, se recreará: %s", e)
        return {}

# --------------------------------------------------------------------------- #
# Construcción del índice en lotes
# --------------------------------------------------------------------------- #

def _procesar_lote(
    paths: list[str],
    vectorstore: Optional[FAISS],
    splitter: RecursiveCharacterTextSplitter,
    embedding_model,
    index_path: str,
) -> tuple[FAISS, list[LCDocument]]:
    documents: list[LCDocument] = []
    for p in paths:
        documents.extend(load_document(p))

    if not documents:
        return vectorstore, []

    splits = splitter.split_documents(documents)
    del documents

    for chunk in splits:
        res = chunk.metadata.get("resolucion", "Sin identificador")
        name = chunk.metadata.get("source", "Sin archivo")
        anio = chunk.metadata.get("anio") or "s/d"
        chunk.page_content = (
            f"[Resolución: {res} | Año: {anio} | Archivo: {name}]\n"
            f"{chunk.page_content}"
        )

    if vectorstore is None:
        vectorstore = FAISS.from_documents(splits, embedding_model)
    else:
        vectorstore.add_documents(splits)

    vectorstore.save_local(index_path)
    return vectorstore, splits

@st.cache_resource(show_spinner="Cargando índice de resoluciones…")
def load_vectorstore():
    BASE_PATH = os.path.dirname(os.path.abspath(__file__))
    DOCS_PATH = os.path.join(BASE_PATH, "RESOLUCIONES_2", "RESOLUCIONES")
    INDEX_PATH = os.path.join(BASE_PATH, "faiss_index")
    ESTADO_PATH = os.path.join(BASE_PATH, "estado_indices.pkl")
    BM25_PATH = os.path.join(BASE_PATH, "bm25_corpus.pkl")

    os.makedirs(DOCS_PATH, exist_ok=True)
    os.makedirs(INDEX_PATH, exist_ok=True)

    estado_prev = cargar_estado(ESTADO_PATH)
    all_paths = (
        glob.glob(f"{DOCS_PATH}/**/*.docx", recursive=True)
        + glob.glob(f"{DOCS_PATH}/**/*.pdf", recursive=True)
    )

    nuevos: list[str] = []
    estado_actual: dict[str, tuple[float, int]] = {}
    for p in all_paths:
        try:
            sig = _file_signature(p)
        except OSError:
            continue
        estado_actual[p] = sig
        if estado_prev.get(p) != sig:
            nuevos.append(p)

    logger.info(
        "Total=%d | nuevos=%d | device=%s | RAM libre=%.1f GB",
        len(all_paths), len(nuevos), DEVICE,
        _ram_disponible_gb() or -1,
    )

    embedding_model = HuggingFaceEmbeddings(
        model_name="BAAI/bge-m3",
        model_kwargs={"device": DEVICE, "trust_remote_code": True},
        encode_kwargs={"normalize_embeddings": True, "batch_size": 16},
    )

    index_file = os.path.join(INDEX_PATH, "index.faiss")
    if os.path.exists(index_file) and estado_prev:
        vectorstore = FAISS.load_local(
            INDEX_PATH, embedding_model, allow_dangerous_deserialization=True,
        )
        logger.info("Índice FAISS cargado desde disco.")
    else:
        vectorstore = None
        logger.info("No hay índice FAISS válido — se creará uno nuevo.")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n--- Tabla", "\n\n", "\n", ". ", " ", ""],
    )

    nuevos_chunks_acum: list[LCDocument] = []
    if nuevos:
        st.warning(
            f"Indexando {len(nuevos)} documentos. "
            f"Esto puede tardar varios minutos. **No cierres la app.** "
            f"Si te quedaste sin RAM antes, cerrá Ollama mientras dura la indexación inicial."
        )
        progress = st.progress(0.0, text="Iniciando indexación…")
        total = len(nuevos)
        procesados = 0

        for i in range(0, total, BATCH_SIZE_DOCS):
            lote = nuevos[i:i + BATCH_SIZE_DOCS]
            progress.progress(
                procesados / total,
                text=f"Lote {i // BATCH_SIZE_DOCS + 1} "
                     f"({procesados}/{total} documentos)",
            )

            vectorstore, chunks_lote = _procesar_lote(
                lote, vectorstore, splitter, embedding_model, INDEX_PATH,
            )
            nuevos_chunks_acum.extend(chunks_lote)
            procesados += len(lote)

            estado_parcial = {p: estado_actual[p] for p in nuevos[:procesados]
                              if p in estado_actual}
            estado_parcial.update(estado_prev)
            guardar_estado(ESTADO_PATH, estado_parcial)

            del chunks_lote
            if procesados % GC_EVERY_N_DOCS == 0:
                gc.collect()

        progress.progress(1.0, text="Indexación completada.")
        guardar_estado(ESTADO_PATH, estado_actual)

    bm25_corpus = _cargar_o_construir_bm25_corpus(
        BM25_PATH, nuevos_chunks_acum, vectorstore,
    )

    return vectorstore, embedding_model, bm25_corpus

def _cargar_o_construir_bm25_corpus(
    bm25_path: str,
    nuevos_chunks: list[LCDocument],
    vectorstore: FAISS,
) -> list[LCDocument]:
    if os.path.exists(bm25_path):
        try:
            with open(bm25_path, "rb") as f:
                corpus = pickle.load(f)
            if nuevos_chunks:
                corpus.extend(nuevos_chunks)
                with open(bm25_path, "wb") as f:
                    pickle.dump(corpus, f)
            logger.info("BM25 corpus cargado: %d chunks", len(corpus))
            return corpus
        except Exception as e:  # noqa: BLE001
            logger.warning("BM25 cache corrupto, se reconstruye: %s", e)

    logger.info("Reconstruyendo BM25 corpus desde FAISS docstore…")
    corpus: list[LCDocument] = []
    try:
        ds = vectorstore.docstore._dict  # noqa: SLF001
        corpus = list(ds.values())
    except Exception as e:  # noqa: BLE001
        logger.error("No se pudo extraer corpus desde FAISS: %s", e)
        corpus = list(nuevos_chunks)

    if corpus:
        with open(bm25_path, "wb") as f:
            pickle.dump(corpus, f)
        logger.info("BM25 corpus guardado: %d chunks", len(corpus))
    gc.collect()
    return corpus

# --------------------------------------------------------------------------- #
# Cross-encoder (re-ranking obligatorio)
# --------------------------------------------------------------------------- #

@st.cache_resource(show_spinner="Cargando re-ranker…")
def load_reranker():
    from sentence_transformers import CrossEncoder
    return CrossEncoder(
        "cross-encoder/mmarco-mMiniLMv2-L12-H384-v1",
        device="cpu",
        max_length=512,
    )

# --------------------------------------------------------------------------- #
# Intención temporal
# --------------------------------------------------------------------------- #

PALABRAS_PRESENTE = re.compile(
    r"\b(actual(?:mente)?|hoy|vigente|vigentes|ahora|"
    r"quién es|quien es|cuál es|cual es|cuáles son|cuales son|"
    r"está\b|esta\b|son\b|es el|es la|es del)\b",
    re.IGNORECASE,
)
PALABRAS_HISTORICAS = re.compile(
    r"\b(en \d{4}|en el \d{4}|hist[oó]ric|antes|anteriormente|"
    r"resoluci[oó]n \w+|fue|era)\b",
    re.IGNORECASE,
)

def intencion_temporal(pregunta: str) -> str:
    if PALABRAS_HISTORICAS.search(pregunta):
        return "historico"
    if PALABRAS_PRESENTE.search(pregunta):
        return "reciente"
    return "neutro"

# --------------------------------------------------------------------------- #
# Detección de identificadores de resolución en la query
# --------------------------------------------------------------------------- #

PATRON_ID_EN_QUERY = re.compile(
    r"(?<![A-Za-z])"
    r"([A-Z]{1,6})"
    r"[\s\-\u2010-\u2015]?"
    r"(\d{1,4})"
    r"\s*[\-\.\/\u2010-\u2015]\s*"
    r"(\d{2,4})"
    r"(?!\d)",
    re.IGNORECASE,
)

_FALSOS_PREFIJOS = {
    "LA", "EL", "DE", "EN", "UN", "MI", "TU", "SU", "LO", "LE",
    "AL", "DEL", "POR", "QUE", "ES", "SE", "ME", "TE", "ESTA", "ESTE",
    "ESA", "ESE", "PARA", "CON", "SIN", "Y", "O", "U", "A",
    "LAS", "LOS", "UNA", "UNAS", "UNOS", "NO", "SI", "ASI",
}

def extraer_ids_de_query(query: str) -> list[str]:
    matches = PATRON_ID_EN_QUERY.findall(query)
    ids: list[str] = []
    for letras, n1, n2 in matches:
        if letras.upper() in _FALSOS_PREFIJOS:
            continue
        ids.append(f"{letras.upper()}-{n1}.{n2}")
    return ids

def _id_matchea(id_query: str, id_doc: str) -> bool:
    def _norm(s: str) -> str:
        s = s.upper().strip()
        s = re.sub(r"[\-\u2010-\u2015]", ".", s)
        s = re.sub(r"[\s/]", ".", s)
        s = re.sub(r"\.+", ".", s)
        return s
    n_q = _norm(id_query)
    n_d = _norm(id_doc)
    return n_q == n_d or n_q in n_d or n_d in n_q

# --------------------------------------------------------------------------- #
# Normalización para BM25
# --------------------------------------------------------------------------- #

def _quitar_acentos(texto: str) -> str:
    nfd = unicodedata.normalize("NFD", texto)
    return "".join(c for c in nfd if unicodedata.category(c) != "Mn")

_TOKEN_RE = re.compile(r"[a-z0-9]+")

def bm25_preprocess(texto: str) -> list[str]:
    """Lowercase + sin acentos, para que 'Cátedra' matchee 'catedra'."""
    texto = _quitar_acentos(texto.lower())
    return _TOKEN_RE.findall(texto)

# --------------------------------------------------------------------------- #
# Retriever híbrido: RRF + cross-encoder + recencia + IDs + expansión
# --------------------------------------------------------------------------- #

class HybridRecencyRetriever:
    RRF_K = 60
    MAX_CHUNKS_POR_RES = 5
    MAX_CHUNKS_DOC_PINEADO = 8

    def __init__(
        self,
        vectorstore: FAISS,
        bm25_corpus: list[LCDocument],
        reranker,
    ):
        self.vectorstore = vectorstore
        self.reranker = reranker
        self.bm25_corpus = bm25_corpus
        self.last_debug: dict = {}

        self.faiss_retriever = vectorstore.as_retriever(
            search_type="mmr",
            search_kwargs={"k": TOP_K_INITIAL, "fetch_k": TOP_K_INITIAL * 2},
        )
        self.bm25 = BM25Retriever.from_documents(
            bm25_corpus, preprocess_func=bm25_preprocess,
        )
        self.bm25.k = TOP_K_INITIAL

        self.por_resolucion: dict[str, list[LCDocument]] = {}
        for chunk in bm25_corpus:
            res = chunk.metadata.get("resolucion", "")
            if res:
                self.por_resolucion.setdefault(res, []).append(chunk)
        logger.info(
            "Índice por resolución: %d resoluciones, %d chunks",
            len(self.por_resolucion), len(bm25_corpus),
        )

    @staticmethod
    def _normaliza(q: str) -> str:
        q = re.sub(r"\bAyC\b", "Algoritmos y Complejidad", q, flags=re.I)
        q = re.sub(r"\bBD\b", "Bases de Datos", q, flags=re.I)
        return q.strip()

    @staticmethod
    def _doc_key(doc: LCDocument) -> str:
        return doc.metadata.get("path", "") + "|" + str(hash(doc.page_content[:200]))

    def _chunks_de_documento(self, ids_query: list[str]) -> list[LCDocument]:
        """Pinning por ID explícito en la pregunta."""
        if not ids_query:
            return []
        encontrados: list[LCDocument] = []
        for id_doc, chunks in self.por_resolucion.items():
            if any(_id_matchea(id_q, id_doc) for id_q in ids_query):
                encontrados.extend(chunks)
        if encontrados:
            logger.info(
                "IDs en query=%s → %d chunks pineados",
                ids_query, len(encontrados),
            )
        return encontrados

    def _rrf_fusion(
        self,
        rankings: list[list[LCDocument]],
    ) -> list[tuple[LCDocument, float]]:
        scores: dict[str, float] = {}
        index: dict[str, LCDocument] = {}
        for ranking in rankings:
            for rank, doc in enumerate(ranking):
                k = self._doc_key(doc)
                scores[k] = scores.get(k, 0.0) + 1.0 / (self.RRF_K + rank)
                if k not in index:
                    index[k] = doc

        return [(index[k], s) for k, s in sorted(scores.items(), key=lambda x: -x[1])]

    def _rerank_cross_encoder(
        self,
        query: str,
        scored: list[tuple[LCDocument, float]],
    ) -> list[tuple[LCDocument, float]]:
        if not scored:
            return scored
        docs = [d for d, _ in scored]
        pairs = [(query, d.page_content) for d in docs]
        ce_scores = self.reranker.predict(
            pairs, batch_size=8, show_progress_bar=False,
        )
        out = []
        for (doc, rrf), ce in zip(scored, ce_scores):
            out.append((doc, float(ce) + 0.1 * rrf))
        return out

    @staticmethod
    def _aplicar_recencia(
        scored: list[tuple[LCDocument, float]],
        intencion: str,
    ) -> list[tuple[LCDocument, float]]:
        # Intención "neutro" aplica un desempate temporal suave: entre
        # candidatos casi empatados gana la resolución más nueva, sin
        # sobrescribir diferencias de relevancia grandes.
        ajustados: list[tuple[LCDocument, float]] = []
        for doc, s in scored:
            anio = doc.metadata.get("anio") or 0
            if intencion == "reciente":
                peso = RECENCY_BOOST
                if anio == 0:
                    ajuste = -0.1
                else:
                    ajuste = -min(1.0, (CURRENT_YEAR - anio) / 10.0)
            elif intencion == "historico":
                peso = RECENCY_BOOST
                ajuste = 0.0
            else:  # neutro → desempate suave
                peso = RECENCY_TIEBREAK_NEUTRO
                if anio == 0:
                    ajuste = -0.05
                else:
                    ajuste = -min(1.0, (CURRENT_YEAR - anio) / 10.0)
            ajustados.append((doc, s + peso * ajuste))
        return ajustados

    @staticmethod
    def _dedupe(
        scored: list[tuple[LCDocument, float]],
        max_por_res: int,
    ) -> list[tuple[LCDocument, float]]:
        vistos: dict[str, int] = {}
        out: list[tuple[LCDocument, float]] = []
        for doc, s in scored:
            res = doc.metadata.get("resolucion", "sin_id")
            if vistos.get(res, 0) >= max_por_res:
                continue
            vistos[res] = vistos.get(res, 0) + 1
            out.append((doc, s))
        return out

    def _expandir_resoluciones(
        self,
        seleccionados: list[LCDocument],
        n_res: int = N_RESOLUCIONES_EXPANDIR,
    ) -> list[LCDocument]:
        """Trae todos los chunks de las top-N resoluciones del ranking,
        para recuperar tablas y listas que el splitter haya cortado."""
        out = list(seleccionados)
        presentes = {self._doc_key(d) for d in out}

        res_top: list[str] = []
        for d in seleccionados:
            r = d.metadata.get("resolucion", "")
            if r and r not in res_top:
                res_top.append(r)
            if len(res_top) >= n_res:
                break

        agregados = 0
        for r in res_top:
            hermanos = self.por_resolucion.get(r, [])[:MAX_CHUNKS_POR_RES_EXPANDIDA]
            for sib in hermanos:
                k = self._doc_key(sib)
                if k not in presentes:
                    out.append(sib)
                    presentes.add(k)
                    agregados += 1
        if agregados:
            logger.info(
                "Expansión: +%d chunks de %s", agregados, res_top,
            )
        return out

    @staticmethod
    def _aplicar_presupuesto(
        docs: list[LCDocument],
        max_chars: int = MAX_CONTEXT_CHARS,
    ) -> list[LCDocument]:
        """Recorta la lista (ordenada por importancia) al presupuesto de
        caracteres, para no exceder num_ctx del LLM."""
        out: list[LCDocument] = []
        total = 0
        for d in docs:
            n = len(d.page_content)
            if out and total + n > max_chars:
                break
            out.append(d)
            total += n
        return out

    @staticmethod
    def _reordenar_anti_lost_middle(docs: list[LCDocument]) -> list[LCDocument]:
        if len(docs) <= 2:
            return docs
        inicio = docs[0::2]
        final = docs[1::2][::-1]
        return inicio + final

    def get_relevant_documents(
        self,
        query: str,
        query_alts: Optional[list[str]] = None,
    ) -> list[LCDocument]:
        """`query` es siempre la pregunta original del usuario. `query_alts`
        son variantes adicionales (reescritura con historial y/o multi-query).
        Se busca con la original y con cada variante, fusionando con RRF."""
        q = self._normaliza(query)
        intencion = intencion_temporal(q)
        variantes = [v for v in (query_alts or []) if v and v.strip()]

        ids_en_query = extraer_ids_de_query(q)
        for v in variantes:
            ids_en_query.extend(extraer_ids_de_query(v))
        ids_en_query = list(dict.fromkeys(ids_en_query))

        chunks_pineados = self._chunks_de_documento(ids_en_query)

        textos_busqueda = [q]
        for v in variantes:
            vn = self._normaliza(v)
            if vn and vn != q and vn not in textos_busqueda:
                textos_busqueda.append(vn)

        rankings: list[list[LCDocument]] = []
        for t in textos_busqueda:
            rankings.append(self.bm25.invoke(t))
            rankings.append(self.faiss_retriever.invoke(t))

        fused = self._rrf_fusion(rankings)
        fused = self._rerank_cross_encoder(q, fused)
        fused = self._aplicar_recencia(fused, intencion)
        fused.sort(key=lambda x: x[1], reverse=True)
        fused = self._dedupe(fused, max_por_res=self.MAX_CHUNKS_POR_RES)

        if chunks_pineados:
            pineados = chunks_pineados[:self.MAX_CHUNKS_DOC_PINEADO]
            claves_pin = {self._doc_key(c) for c in pineados}
            extras = [d for d, _ in fused if self._doc_key(d) not in claves_pin]
            slots_restantes = max(0, TOP_K_FINAL - len(pineados))
            combinado = pineados + extras[:slots_restantes]
        else:
            combinado = [d for d, _ in fused[:TOP_K_FINAL]]

        combinado = self._expandir_resoluciones(combinado)
        combinado = self._aplicar_presupuesto(combinado)
        final = self._reordenar_anti_lost_middle(combinado)

        self.last_debug = {
            "query_original": query,
            "variantes_busqueda": textos_busqueda[1:],
            "intencion": intencion,
            "ids_detectados": ids_en_query,
            "chunks_pineados": len(chunks_pineados),
            "chunks_finales": len(final),
            "chars_contexto": sum(len(d.page_content) for d in final),
            "ranking": [
                {
                    "resolucion": d.metadata.get("resolucion"),
                    "anio": d.metadata.get("anio"),
                    "archivo": d.metadata.get("source"),
                }
                for d in final
            ],
        }
        logger.info(
            "Q=%r | variantes=%d | intencion=%s | ids=%s | pineados=%d | "
            "finales=%d | chars=%d | años=%s",
            query, len(textos_busqueda) - 1, intencion, ids_en_query,
            len(chunks_pineados), len(final),
            self.last_debug["chars_contexto"],
            [d.metadata.get("anio") for d in final],
        )
        return final

# --------------------------------------------------------------------------- #
# Sidebar (diagnóstico)
# --------------------------------------------------------------------------- #

with st.sidebar:
    st.subheader("⚙️ Opciones")
    mostrar_fuentes = st.checkbox("Mostrar fuentes recuperadas", value=True)
    modo_debug = st.checkbox(
        "🔬 Modo debug",
        value=False,
        help="Muestra la query usada para buscar, IDs detectados y ranking final.",
    )

    st.divider()
    st.subheader("🩺 Diagnóstico")
    st.caption(f"Device embeddings: **{DEVICE}**")
    ram = _ram_disponible_gb()
    if ram is not None:
        color = "🟢" if ram > 4 else "🟡" if ram > 2 else "🔴"
        st.caption(f"RAM libre: {color} **{ram:.1f} GB**")
    if _ollama_running():
        st.caption("Ollama: 🟢 corriendo")
    else:
        st.caption("Ollama: 🔴 no detectado")
    st.caption(f"Año actual: **{CURRENT_YEAR}**")

    if st.button("🗑️ Limpiar conversación"):
        st.session_state.history = []
        st.rerun()

# --------------------------------------------------------------------------- #
# Carga de recursos
# --------------------------------------------------------------------------- #

vectorstore, embedding, bm25_corpus = load_vectorstore()
if vectorstore is None or not bm25_corpus:
    st.error("No se pudo cargar el índice o el corpus BM25.")
    st.stop()

reranker = load_reranker()
hybrid = HybridRecencyRetriever(vectorstore, bm25_corpus, reranker)

# --------------------------------------------------------------------------- #
# LLM y cadena RAG manual
# --------------------------------------------------------------------------- #

from langchain_ollama import OllamaLLM

llm = OllamaLLM(
    model="gemma3:12b",
    temperature=0,
    num_ctx=8192,
    num_predict=700,
)

# Reescritura de la pregunta con historial, en español, preservando
# identificadores de resolución textualmente.
CONDENSE_TEMPLATE = """\
Dada la siguiente conversación y una pregunta de seguimiento, reescribí la \
pregunta de seguimiento como una pregunta independiente y autocontenida, \
EN ESPAÑOL.

REGLAS:
- Si la pregunta menciona números o identificadores de resolución \
(ej. "CECC-1/95", "CSU-330-19"), copialos EXACTAMENTE igual, sin modificarlos.
- Si la pregunta ya es autocontenida, devolvela tal cual, sin cambios.
- Respondé ÚNICAMENTE con la pregunta reescrita, sin explicaciones ni comillas.

CONVERSACIÓN:
{historial}

PREGUNTA DE SEGUIMIENTO:
{pregunta}

PREGUNTA INDEPENDIENTE:"""

def condensar_pregunta(
    query: str,
    chat_history: list[tuple[str, str]],
) -> Optional[str]:
    """Genera la versión autocontenida de la pregunta usando el historial.
    Devuelve None si no hay historial o si la reescritura falla la
    validación (en ese caso se busca solo con la pregunta original)."""
    if not chat_history:
        return None
    historial = "\n".join(
        f"Usuario: {q}\nAsistente: {a[:400]}" for q, a in chat_history
    )
    try:
        out = llm.invoke(
            CONDENSE_TEMPLATE.format(historial=historial, pregunta=query)
        ).strip().strip('"')
    except Exception as e:  # noqa: BLE001
        logger.warning("Falló la reescritura de pregunta: %s", e)
        return None

    if not out or len(out) > MAX_LEN_PREGUNTA_REESCRITA:
        return None
    ids_orig = extraer_ids_de_query(query)
    ids_nueva = extraer_ids_de_query(out)
    if ids_orig and not all(
        any(_id_matchea(io, ino) for ino in ids_nueva) for io in ids_orig
    ):
        logger.warning(
            "Reescritura descartada: perdió IDs %s → %r", ids_orig, out,
        )
        return None
    return out if out != query else None


# Multi-query expansion: reformulaciones en el vocabulario probable del
# documento (correlativa, requisito, excepción, etc.). Genérico: no
# contiene ninguna respuesta ni materia hardcodeada.
MULTIQUERY_TEMPLATE = """\
Sos un experto en resoluciones académicas. A partir de la siguiente consulta, \
generá {n} reformulaciones distintas que sirvan para BUSCAR el documento que \
la responde. Las reformulaciones deben usar el vocabulario que probablemente \
aparezca en el texto oficial.

REGLAS:
- Si la consulta tiene una negación ("sin aprobar", "sin haber rendido", "sin \
necesidad de"), reformulá también desde la perspectiva del documento: \
"se elimina la correlativa", "deja de ser requisito/correlativa", \
"ya no se exige", "se suprime el requisito", "excepción".
- Usá sinónimos institucionales: correlativa, requisito, prerrequisito, \
habilitación, excepción, modificación del plan de estudios.
- Conservá EXACTAMENTE los nombres de materias, personas, fechas e \
identificadores que aparezcan EN LA CONSULTA.
- PROHIBIDO usar marcadores de posición entre corchetes (como "[Materia]", \
"[Nombre]", "[Materia a habilitar]"). Si un dato NO aparece en la consulta \
porque es justamente lo que se pregunta, NO lo inventes ni lo marques: dejá la \
reformulación en términos generales (ej. "materia habilitada a cursar sin \
aprobar Sistemas Operativos", "excepción a la correlativa Sistemas Operativos").
- PROHIBIDO desarrollar, traducir o adivinar el significado de SIGLAS y \
ACRÓNIMOS (ej. "DCIC", "ISI", "LCC", "IAW", "PROMINF"). Copiá la sigla TAL CUAL \
aparece en la consulta. NUNCA la reemplaces por una expansión inventada \
(ej. NO escribas "Facultad de Ciencias Informáticas" ni "Distrito Cibernético" \
en lugar de "DCIC"): una expansión equivocada arruina la búsqueda.
- Una reformulación por línea, sin numerar, sin comillas, sin explicaciones.

CONSULTA:
{pregunta}

REFORMULACIONES:"""

def _siglas_en(texto: str) -> set[str]:
    """Siglas/acrónimos (2+ letras mayúsculas) presentes en el texto."""
    return set(re.findall(r"\b[A-Z]{2,}\b", texto))

def generar_variantes_query(query: str, n: int = N_VARIANTES_QUERY) -> list[str]:
    """Pide al LLM N reformulaciones de la consulta. Cada una se usa como
    búsqueda adicional (BM25 + FAISS) y se fusiona con RRF. Si el LLM falla
    o devuelve basura, se devuelve [] y el sistema sigue funcionando igual
    que sin multi-query."""
    try:
        out = llm.invoke(MULTIQUERY_TEMPLATE.format(n=n, pregunta=query))
    except Exception as e:  # noqa: BLE001
        logger.warning("Falló multi-query: %s", e)
        return []

    siglas_query = _siglas_en(query)
    variantes: list[str] = []
    for linea in out.splitlines():
        v = linea.strip().lstrip("-•*0123456789.) ").strip().strip('"')
        if not v or len(v) > MAX_LEN_VARIANTE:
            continue
        if "[" in v or "]" in v:
            continue
        if siglas_query and not (siglas_query & _siglas_en(v)):
            logger.info("Variante descartada (perdió sigla %s): %r", siglas_query, v)
            continue
        if v.lower() == query.strip().lower():
            continue
        if v not in variantes:
            variantes.append(v)
    variantes = variantes[:n]
    if variantes:
        logger.info("Multi-query generó %d variantes: %s", len(variantes), variantes)
    return variantes


PROMPT_TEMPLATE = """\
Eres un asistente experto en analizar documentos y resoluciones académicas institucionales. Respondé SIEMPRE en español de forma directa y clara.

REGLAS ESTRICTAS:
1. Usá únicamente la información del CONTEXTO. Si no alcanza, respondé exactamente:
   "No se encontró información suficiente en las resoluciones disponibles."
2. PRIORIDAD ABSOLUTA DE RESOLUCIÓN: Si el usuario menciona un número o identificador específico (ej. "CECC-1/95"), buscá la respuesta EXCLUSIVAMENTE en esos chunks.
3. ROLES Y CARGOS DOCENTES (¡MUY IMPORTANTE!): Prestá extrema atención a la jerarquía de cargos.
   - Si el usuario pregunta por el "Profesor" (o Titular/Adjunto) de una materia, NO respondas con el nombre de un "Ayudante", "Asistente" o "JTP".
   - Analizá las estructuras "Cargo: [Rol] | Nombre: [Persona]" o el texto circundante para emparejar a la persona correcta con el cargo exacto que se está consultando.
4. CONTRADICCIONES Y VIGENCIA: Si hay información contradictoria o varias resoluciones sobre el mismo tema:
   - Una resolución que DEROGA, SUSPENDE, DEJA SIN EFECTO o RECTIFICA a otra hace que la disposición derogada/suspendida YA NO ESTÉ VIGENTE. NO des como respuesta una disposición que fue derogada o suspendida: si en el contexto hay una resolución que deroga/suspende una excepción o correlativa, esa excepción ya no aplica.
   - La respuesta vigente es la disposición que NO fue derogada, o la rectificación/versión más reciente que sigue en vigor.
   - Si el usuario NO especificó una resolución concreta, priorizá la información VIGENTE más reciente y mencioná brevemente que existieron versiones o excepciones anteriores (y cuáles fueron derogadas).
5. CITA DE FUENTES: Mencioná explícitamente el número de resolución y el año de donde extraés cada dato.
6. Al final, agregá la referencia en el formato:
   [Resolución: <id> | Año: <año> | Archivo: <nombre>]

CONTEXTO:
{context}

PREGUNTA:
{question}

RESPUESTA:"""

def responder(query: str, chat_history: list[tuple[str, str]]):
    """Cadena RAG manual:
    1. Reescribe la pregunta con historial si hay turnos previos.
    2. Genera variantes multi-query sobre la pregunta autocontenida.
    3. Recupera con la original + reescrita + variantes (fusión RRF).
    4. Arma el contexto y consulta al LLM una sola vez.
    """
    standalone = condensar_pregunta(query, chat_history)
    base = standalone or query

    variantes = generar_variantes_query(base)
    query_alts = ([standalone] if standalone else []) + variantes

    docs = hybrid.get_relevant_documents(query, query_alts=query_alts)

    context = "\n\n".join(d.page_content for d in docs)
    prompt = PROMPT_TEMPLATE.format(context=context, question=base)
    respuesta = llm.invoke(prompt).strip()
    return respuesta, docs

# --------------------------------------------------------------------------- #
# UI principal
# --------------------------------------------------------------------------- #

if "history" not in st.session_state:
    st.session_state.history = []

query = st.text_input("Escribí tu pregunta:")

if query:
    with st.spinner("Buscando respuesta…"):
        history_recent = st.session_state.history[-MAX_HISTORY_TURNS:]
        chat_history = [(qa["question"], qa["answer"]) for qa in history_recent]

        respuesta, fuentes = responder(query, chat_history)

        st.session_state.history.append({
            "question": query,
            "answer": respuesta,
            "fuentes": [
                {
                    "resolucion": d.metadata.get("resolucion"),
                    "anio": d.metadata.get("anio"),
                    "archivo": d.metadata.get("source"),
                }
                for d in fuentes
            ],
            "debug": dict(hybrid.last_debug),
        })

for qa in reversed(st.session_state.history):
    st.markdown(f"**Pregunta:** {qa['question']}")
    st.markdown(f"**Respuesta:** {qa['answer']}")
    if mostrar_fuentes and qa.get("fuentes"):
        with st.expander("📚 Fuentes recuperadas"):
            for f in qa["fuentes"]:
                st.markdown(
                    f"- **{f['resolucion']}** "
                    f"({f['anio'] or 's/d'}) — `{f['archivo']}`"
                )
    if modo_debug and qa.get("debug"):
        with st.expander("🔬 Debug de retrieval"):
            st.json(qa["debug"])
    st.markdown("---")
